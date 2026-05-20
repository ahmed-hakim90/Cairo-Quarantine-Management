#!/usr/bin/env python3
"""Build docs/HANDOVER-AR.docx from docs/HANDOVER-AR.md with embedded screenshots."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
MD_PATH = DOCS / "HANDOVER-AR.md"
OUT_PATH = DOCS / "HANDOVER-AR.docx"
SCREENSHOTS_DIR = DOCS / "handover" / "screenshots"

IMAGE_RE = re.compile(r"!\[[^\]]*\]\(([^)]+)\)")
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")
CODE_INLINE_RE = re.compile(r"`([^`]+)`")
IMG_IN_CELL_RE = re.compile(r"!\[[^\]]*\]\([^)]+\)")

MAX_IMAGE_WIDTH = Inches(6.2)


def set_rtl_paragraph(paragraph) -> None:
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = p_pr.makeelement(qn("w:bidi"))
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def add_formatted_runs(paragraph, text: str, *, bold_default: bool = False) -> None:
    """Add runs with **bold** and `code` markers."""
    pos = 0
    combined = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    for match in combined.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            run.bold = bold_default
        chunk = match.group(0)
        if chunk.startswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.bold = bold_default


def resolve_image(rel_path: str) -> Path | None:
    rel = rel_path.strip()
    candidates = [
        DOCS / rel,
        ROOT / rel,
        SCREENSHOTS_DIR / Path(rel).name,
    ]
    for path in candidates:
        if path.is_file():
            return path
    return None


def add_image_block(doc: Document, rel_path: str, caption: str | None = None) -> None:
    path = resolve_image(rel_path)
    if not path:
        p = doc.add_paragraph()
        set_rtl_paragraph(p)
        add_formatted_runs(p, f"[صورة غير موجودة: {rel_path}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run()
    try:
        run.add_picture(str(path), width=MAX_IMAGE_WIDTH)
    except Exception as exc:  # noqa: BLE001
        p2 = doc.add_paragraph()
        set_rtl_paragraph(p2)
        add_formatted_runs(p2, f"[تعذر إدراج الصورة: {path.name} — {exc}]")
        return
    if caption:
        cap = doc.add_paragraph()
        set_rtl_paragraph(cap)
        run_cap = cap.add_run(caption)
        run_cap.italic = True
        run_cap.font.size = Pt(10)


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [cell.strip() for cell in line.split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|[\s\-:|]+\|\s*$", line.strip()))


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            cell_text = row[c_idx] if c_idx < len(row) else ""
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            img_match = IMG_IN_CELL_RE.search(cell_text)
            if img_match:
                rel = IMAGE_RE.search(cell_text)
                if rel:
                    path = resolve_image(rel.group(1))
                    if path:
                        p = cell.paragraphs[0]
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run()
                        try:
                            run.add_picture(str(path), width=Inches(2.8))
                        except Exception:
                            p.text = path.name
                    cell_text = IMG_IN_CELL_RE.sub("", cell_text).strip()
            if cell_text:
                p = cell.paragraphs[0]
                set_rtl_paragraph(p)
                add_formatted_runs(p, cell_text)


def convert_markdown_to_docx(md_text: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_lang = ""
    code_lines: list[str] = []
    table_rows: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_rows
        if table_rows:
            add_table(doc, table_rows)
            table_rows = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if in_code:
            if stripped.startswith("```"):
                p = doc.add_paragraph()
                set_rtl_paragraph(p)
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                if code_lang == "mermaid":
                    note = doc.add_paragraph()
                    set_rtl_paragraph(note)
                    nr = note.add_run("(مخطط Mermaid — راجع النسخة الإلكترونية أو README)")
                    nr.italic = True
                    nr.font.size = Pt(9)
                in_code = False
                code_lines = []
                code_lang = ""
            else:
                code_lines.append(line)
            i += 1
            continue

        if stripped.startswith("```"):
            flush_table()
            in_code = True
            code_lang = stripped[3:].strip()
            i += 1
            continue

        if stripped.startswith("|") and not is_table_separator(stripped):
            flush_table()
            table_rows.append(parse_table_row(stripped))
            i += 1
            continue

        if is_table_separator(stripped):
            i += 1
            continue

        flush_table()

        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        img_only = IMAGE_RE.fullmatch(stripped)
        if img_only:
            add_image_block(doc, img_only.group(1))
            i += 1
            continue

        for img_match in IMAGE_RE.finditer(stripped):
            if img_match.group(0) == stripped:
                add_image_block(doc, img_match.group(1))
                i += 1
                continue

        if stripped.startswith("# "):
            h = doc.add_heading(stripped[2:].strip(), level=0)
            set_rtl_paragraph(h)
            i += 1
            continue
        if stripped.startswith("## "):
            h = doc.add_heading(stripped[3:].strip(), level=1)
            set_rtl_paragraph(h)
            i += 1
            continue
        if stripped.startswith("### "):
            h = doc.add_heading(stripped[4:].strip(), level=2)
            set_rtl_paragraph(h)
            i += 1
            continue

        if stripped.startswith("> "):
            p = doc.add_paragraph()
            set_rtl_paragraph(p)
            add_formatted_runs(p, stripped[2:])
            i += 1
            continue

        if re.match(r"^[-*] ", stripped):
            p = doc.add_paragraph(style="List Bullet")
            set_rtl_paragraph(p)
            add_formatted_runs(p, stripped[2:])
            i += 1
            continue

        if re.match(r"^\d+\. ", stripped):
            p = doc.add_paragraph(style="List Number")
            set_rtl_paragraph(p)
            add_formatted_runs(p, re.sub(r"^\d+\.\s*", "", stripped))
            i += 1
            continue

        if stripped.startswith("- [ ] ") or stripped.startswith("- [x] "):
            p = doc.add_paragraph()
            set_rtl_paragraph(p)
            checked = stripped.startswith("- [x] ")
            add_formatted_runs(p, ("☑ " if checked else "☐ ") + stripped[6:].strip())
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        # Inline images in paragraph
        if IMAGE_RE.search(stripped):
            parts = IMAGE_RE.split(stripped)
            imgs = IMAGE_RE.findall(stripped)
            idx_img = 0
            for part_idx, part in enumerate(parts):
                if part.strip():
                    p = doc.add_paragraph()
                    set_rtl_paragraph(p)
                    text = LINK_RE.sub(r"\1", part)
                    add_formatted_runs(p, text)
                if idx_img < len(imgs) and part_idx < len(parts) - 1:
                    add_image_block(doc, imgs[idx_img], caption=Path(imgs[idx_img]).name)
                    idx_img += 1
            i += 1
            continue

        p = doc.add_paragraph()
        set_rtl_paragraph(p)
        text = stripped
        text = LINK_RE.sub(r"\1", text)
        add_formatted_runs(p, text)
        i += 1

    flush_table()
    return doc


def append_screenshots_gallery(doc: Document) -> None:
    doc.add_page_break()
    h = doc.add_heading("ملحق: معرض لقطات الشاشة (كامل)", level=1)
    set_rtl_paragraph(h)
    intro = doc.add_paragraph()
    set_rtl_paragraph(intro)
    add_formatted_runs(
        intro,
        "جميع الصور المستخدمة في المستند — مرتبة حسب اسم الملف.",
    )
    for path in sorted(SCREENSHOTS_DIR.glob("*.png")):
        title = doc.add_paragraph()
        set_rtl_paragraph(title)
        run = title.add_run(path.stem.replace("-", " "))
        run.bold = True
        add_image_block(doc, f"handover/screenshots/{path.name}", caption=path.name)


def main() -> int:
    if not MD_PATH.is_file():
        print(f"Missing {MD_PATH}", file=sys.stderr)
        return 1
    md_text = MD_PATH.read_text(encoding="utf-8")
    doc = convert_markdown_to_docx(md_text)
    append_screenshots_gallery(doc)
    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")
    print(f"Screenshots: {len(list(SCREENSHOTS_DIR.glob('*.png')))} PNG files")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
